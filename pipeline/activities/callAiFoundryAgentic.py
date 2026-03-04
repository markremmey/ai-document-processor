import azure.durable_functions as df
import asyncio
import logging
import os
import io
from typing import Annotated
from agent_framework.azure import AzureAIAgentClient
from azure.identity.aio import DefaultAzureCredential as AsyncDefaultAzureCredential
from pydantic import Field
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipelineUtils.blob_functions import write_to_blob
from configuration import Configuration

name = "callAiFoundryAgentic"
bp = df.Blueprint()

config = Configuration()
FINAL_OUTPUT_CONTAINER = config.get_value("FINAL_OUTPUT_CONTAINER")
AZURE_AI_PROJECT_ENDPOINT = config.get_value("AZURE_AI_PROJECT_ENDPOINT")
OPENAI_MODEL = config.get_value("OPENAI_MODEL")

# Global variable to store context for the tool
_current_context = {}


def create_word_document(
    title: Annotated[str, Field(description="The title of the Word document.")],
    summary: Annotated[str, Field(description="A brief summary or executive overview to include at the beginning.")],
    content: Annotated[str, Field(description="The main content/body of the document. Use newlines to separate paragraphs.")],
) -> str:
    """
    Creates a Word document with the specified title, summary, and content,
    then uploads it to Azure Blob Storage.
    """
    try:
        blob_name = _current_context.get('blob_name', 'output')
        # Remove original extension and add .docx
        base_name = blob_name.rsplit('.', 1)[0] if '.' in blob_name else blob_name
        output_blob_name = f"{base_name}_summary.docx"

        # Create a new Word document
        doc = Document()

        # Add title
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.style = 'Body Text'

        # Add main content section
        doc.add_heading('Document Content', level=1)

        # Split content by newlines and add as paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style = 'Body Text'

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload to blob storage
        write_to_blob(FINAL_OUTPUT_CONTAINER, output_blob_name, buffer.getvalue())

        logging.info(f"Word document created and uploaded: {output_blob_name}")
        return f"Successfully created Word document: {output_blob_name}"

    except Exception as e:
        logging.error(f"Error creating Word document: {e}")
        return f"Error creating Word document: {str(e)}"


AGENT_INSTRUCTIONS = """You are a document processing assistant. Your task is to analyze the provided document text and create a well-formatted Word document summary.

When processing the document:
1. Identify the main topics and key points
2. Create a clear, descriptive title for the document
3. Write a brief executive summary (2-3 sentences)
4. Organize the main content in a logical, readable format

You MUST use the create_word_document tool to generate the output document. The tool will save the document to blob storage automatically.

Always call the create_word_document tool with:
- title: A descriptive title for the document
- summary: A brief executive summary
- content: The main organized content (use newlines to separate paragraphs)
"""


def _get_async_credential() -> AsyncDefaultAzureCredential:
    """
    Creates an async Azure credential matching the Configuration pattern.
    """
    tenant_id = os.environ.get('AZURE_TENANT_ID', "*")

    if os.environ.get("AZURE_FUNCTIONS_ENVIRONMENT") == "Development":
        return AsyncDefaultAzureCredential(
            additionally_allowed_tenants=tenant_id,
            exclude_environment_credential=True,
            exclude_managed_identity_credential=True,
            exclude_cli_credential=False,
            exclude_powershell_credential=False,
            exclude_shared_token_cache_credential=True,
            exclude_developer_cli_credential=False,
            exclude_interactive_browser_credential=True
        )
    else:
        return AsyncDefaultAzureCredential(
            additionally_allowed_tenants=tenant_id,
            exclude_environment_credential=True,
            exclude_managed_identity_credential=False,
            exclude_cli_credential=True,
            exclude_powershell_credential=True,
            exclude_shared_token_cache_credential=True,
            exclude_developer_cli_credential=True,
            exclude_interactive_browser_credential=True
        )


async def run_agent(document_text: str, blob_name: str) -> str:
    """
    Runs the Azure AI Foundry agent to process the document and create a Word document.

    Args:
        document_text (str): The extracted document text to process.
        blob_name (str): The original blob name for naming the output file.

    Returns:
        str: The result message from the agent.
    """
    # Set context for the tool to access
    _current_context['blob_name'] = blob_name

    async with (
        _get_async_credential() as credential,
        AzureAIAgentClient(
            project_endpoint=AZURE_AI_PROJECT_ENDPOINT,
            model_deployment_name=OPENAI_MODEL,
            async_credential=credential,
            agent_name="DocumentProcessor"
        ).as_agent(
            instructions=AGENT_INSTRUCTIONS,
            tools=create_word_document
        ) as agent,
    ):
        prompt = f"Please analyze the following document text and create a Word document summary:\n\n{document_text}"
        result = await agent.run(prompt)
        return result.text


@bp.function_name(name)
@bp.activity_trigger(input_name="inputData")
def run(inputData: dict):
    """
    Calls the Azure AI Foundry Agentic service to process a document and create a Word document.

    Args:
        inputData (dict): Dictionary containing:
            - text_result (str): The extracted document text to process.
            - instance_id (str): The instance ID for logging purposes.
            - blob_name (str, optional): The original blob name for naming output.

    Returns:
        str: The result from the Azure AI Foundry agent.
    """
    try:
        text_result = inputData.get('text_result', '')
        instance_id = inputData.get('instance_id')
        blob_name = inputData.get('blob_name', f'document_{instance_id}')

        logging.info(f"callAiFoundryAgentic.py: Processing document for instance {instance_id}")

        # Run the async agent function
        response = asyncio.run(run_agent(text_result, blob_name))

        logging.info(f"callAiFoundryAgentic.py: Document processed for instance {instance_id}")

        return response

    except Exception as e:
        logging.error(f"Error processing callAiFoundryAgentic for instance {instance_id}: {e}")
        raise
