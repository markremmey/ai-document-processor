import io
import logging
from typing import Annotated
from pydantic import Field
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipelineUtils.blob_functions import write_to_blob
from configuration import Configuration

config = Configuration()

# Global context for passing runtime data to tool functions
# Set blob_name before invoking the agent
agent_context = {}


def create_word_document(
    title: Annotated[str, Field(description="The title of the Word document.")],
    summary: Annotated[str, Field(description="A brief summary or executive overview to include at the beginning.")],
    content: Annotated[str, Field(description="The main content/body of the document. Use newlines to separate paragraphs.")],
) -> str:
    """
    Creates a Word document with the specified title, summary, and content,
    then uploads it to Azure Blob Storage.
    """
    try:
        # Load config at runtime to avoid import-time failures
        final_output_container = config.get_value("FINAL_OUTPUT_CONTAINER")
        blob_name = agent_context.get('blob_name', 'output')
        # Extract just the filename (strip any container path like "bronze/")
        filename = blob_name.split('/')[-1]
        # Remove original extension and add .docx
        base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
        output_blob_name = f"{base_name}.docx"

        # Create a new Word document
        doc = Document()

        # Add title
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.style = 'Body Text'

        # Add main content section
        doc.add_heading('Document Content', level=1)

        # Split content by newlines and add as paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style = 'Body Text'

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload to blob storage
        write_to_blob(final_output_container, output_blob_name, buffer.getvalue())

        logging.info(f"Word document created and uploaded: {output_blob_name}")
        return f"Successfully created Word document: {output_blob_name}"

    except Exception as e:
        logging.error(f"Error creating Word document: {e}")
        return f"Error creating Word document: {str(e)}"
