"""
Word Document Agent Tool

This agent tool creates Word documents and uploads them to Azure Blob Storage.

Usage:
    from pipelineUtils.agents import create_word_document, set_word_doc_context
    
    # Set context before running agent
    set_word_doc_context(blob_name="my_document.pdf", container="output-container")
    
    # Pass create_word_document as a tool to the agent
    agent = client.as_agent(instructions=..., tools=create_word_document)
"""

import io
import logging
from typing import Annotated

from pydantic import Field
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Context storage for the tool
_word_doc_context = {}

# Lazy-loaded config and container
_config = None
_final_output_container = None


def _get_config():
    """Lazy load configuration to avoid module-level initialization issues."""
    global _config, _final_output_container
    if _config is None:
        from configuration import Configuration
        _config = Configuration()
        _final_output_container = _config.get_value("FINAL_OUTPUT_CONTAINER")
    return _config, _final_output_container


def set_word_doc_context(blob_name: str, container: str = None) -> None:
    """
    Set context for the word document tool before running the agent.
    
    Args:
        blob_name: The original blob name (used for naming the output file)
        container: Optional container override (defaults to FINAL_OUTPUT_CONTAINER)
    """
    _word_doc_context['blob_name'] = blob_name
    _word_doc_context['container'] = container


def create_word_document(
    title: Annotated[str, Field(description="The title of the Word document.")],
    summary: Annotated[str, Field(description="A brief summary or executive overview to include at the beginning.")],
    content: Annotated[str, Field(description="The main content/body of the document. Use newlines to separate paragraphs.")],
) -> str:
    """
    Creates a Word document with the specified title, summary, and content,
    then uploads it to Azure Blob Storage.
    """
    from pipelineUtils.blob_functions import write_to_blob
    
    try:
        _, default_container = _get_config()
        
        blob_name = _word_doc_context.get('blob_name', 'output')
        container = _word_doc_context.get('container') or default_container
        
        # Remove original extension and add .docx
        base_name = blob_name.rsplit('.', 1)[0] if '.' in blob_name else blob_name
        output_blob_name = f"{base_name}_summary.docx"

        # Create a new Word document
        doc = Document()

        # Add title
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.style = 'Body Text'

        # Add main content section
        doc.add_heading('Document Content', level=1)

        # Split content by newlines and add as paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style = 'Body Text'

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload to blob storage
        write_to_blob(container, output_blob_name, buffer.getvalue())

        logging.info(f"Word document created and uploaded: {output_blob_name}")
        return f"Successfully created Word document: {output_blob_name}"

    except Exception as e:
        logging.error(f"Error creating Word document: {e}")
        return f"Error creating Word document: {str(e)}"
